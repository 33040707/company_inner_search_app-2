import os
import glob
import base64
import fitz          # PyMuPDF
import openai
import pandas as pd
from docx import Document
from dotenv import load_dotenv

# .envファイルからAPIキーを読み込む
load_dotenv()
client = openai.Client(api_key=os.getenv("OPENAI_API_KEY"))

# データが格納されているフォルダ
DATA_FOLDER = "data"

def process_pdf(file_path):
    """PDFを画像化し、GPT-4o Visionで読み取る"""
    doc = fitz.open(file_path)
    full_text = ""
    for page_num, page in enumerate(doc):
        print(f"   ... ページ {page_num + 1}/{len(doc)} をAIで読み取り中 ...")
        # 高画質(dpi=400)で画像化
        pix = page.get_pixmap(dpi=400)
        img_bytes = pix.tobytes("jpeg")
        base64_image = base64.b64encode(img_bytes).decode('utf-8')
        
        # GPT-4oの視覚機能でテキスト化
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": "この画像は建設・設計などの業務に関する価格表や仕様書などの社内資料です。表の構造や数値を正確にマークダウン形式で書き起こしてください。"},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{base64_image}",
                                "detail": "high"
                            },
                        },
                    ],
                }
            ],
            max_tokens=3000,
            temperature=0.0,
        )
        full_text += response.choices[0].message.content + "\n\n"
    return full_text


def process_docx(file_path):
    """Wordファイルからテキストと表データを抽出する"""
    doc = Document(file_path)
    full_text = ""
    
    # 段落（通常の文章）の抽出
    for para in doc.paragraphs:
        if para.text.strip():
            full_text += para.text + "\n"
            
    full_text += "\n"
    
    # 表（テーブル）の抽出
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            full_text += " | ".join(row_data) + "\n"
        full_text += "\n"
        
    return full_text


def process_xlsx(file_path):
    """Excelファイルから全シートのデータを抽出する"""
    xls = pd.read_excel(file_path, sheet_name=None, engine='openpyxl')
    full_text = ""
    
    for sheet_name, df in xls.items():
        full_text += f"## シート名: {sheet_name}\n\n"
        # 欠損値を空文字にし、CSV形式でテキスト化することでAIが構造を理解しやすくする
        full_text += df.fillna("").to_csv(index=False) + "\n\n"
        
    return full_text


def convert_all_docs_to_text():
    # PDF, DOCX, XLSX ファイルをすべて検索
    target_files = []
    for ext in ["*.pdf", "*.docx", "*.xlsx"]:
        target_files.extend(glob.glob(os.path.join(DATA_FOLDER, ext)))
    
    if not target_files:
        print("⚠️ dataフォルダ内に変換対象のファイルが見つかりません。")
        return

    print(f"📄 {len(target_files)}件のファイルが見つかりました。変換処理を開始します...\n")

    for file_path in target_files:
        file_name = os.path.basename(file_path)
        base_name = os.path.splitext(file_name)[0]
        ext = os.path.splitext(file_name)[1].lower()
        
        output_txt_path = os.path.join(DATA_FOLDER, f"{base_name}.txt")
        
        # すでに同名のテキストファイルが存在する場合はスキップ
        if os.path.exists(output_txt_path):
            print(f"⏭️ スキップ: {output_txt_path} は既に変換済みです。")
            continue
            
        print(f"🔄 変換を開始します: {file_name}")
        try:
            if ext == ".pdf":
                text_data = process_pdf(file_path)
            elif ext == ".docx":
                text_data = process_docx(file_path)
            elif ext == ".xlsx":
                text_data = process_xlsx(file_path)
            else:
                continue

            # テキストファイルとして保存
            with open(output_txt_path, "w", encoding="utf-8") as f:
                f.write(text_data)
            
            print(f"✅ 変換完了！ {output_txt_path} に保存しました。\n")
            
        except Exception as e:
            print(f"❌ エラーが発生しました ({file_name}): {e}\n")

    print("🎉 すべての処理が完了しました。")


if __name__ == "__main__":
    convert_all_docs_to_text()